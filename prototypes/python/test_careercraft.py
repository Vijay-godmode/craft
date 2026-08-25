"""Integration tests for the local CareerCraft workflow."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import tempfile
import unittest
import json
from unittest.mock import patch

from docx import Document

from app import create_app
from ats_engine import extract_requirements
from job_discovery import classify_qa_role, discover_qa_jobs, matches_market
from workspace_assistant import _resume_chat_reply


SAMPLE_PROFILE = {
    "full_name": "Ava Tester",
    "headline": "QA Automation Engineer",
    "email": "ava@example.com",
    "phone": "+91 99999 00000",
    "location": "Bengaluru, India",
    "linkedin_url": "https://www.linkedin.com/in/ava-tester",
    "summary": "QA engineer with hands-on experience creating reliable web and API test coverage.",
    "skills": ["Selenium", "API Testing", "Postman", "Python", "SQL", "Jira", "Agile"],
    "experience": [
        {
            "title": "QA Engineer",
            "company": "Example Labs",
            "location": "Bengaluru",
            "start_date": "Jan 2023",
            "end_date": "",
            "current": True,
            "bullets": [
                "Created Selenium regression tests and API checks using Postman.",
                "Partnered with engineers in Agile delivery and documented defects in Jira.",
            ],
        }
    ],
    "education": [{"degree": "B.Tech, Computer Science", "school": "Example University", "location": "", "graduation": "2022"}],
}

JOB_DESCRIPTION = """
We are hiring a QA Automation Engineer. Required: Selenium, API testing, SQL,
Python, Jira, regression testing, and Agile delivery. Nice to have: Playwright
and CI/CD. You will design test cases and collaborate with engineers.
"""


class CareerCraftTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        database = str(Path(self.temp_dir.name) / "career.db")
        app = create_app({"TESTING": True, "DATABASE": database, "SECRET_KEY": "test-secret"})
        self.client = app.test_client()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_profile_analysis_and_docx_export(self) -> None:
        response = self.client.put("/api/profile", json=SAMPLE_PROFILE)
        self.assertEqual(response.status_code, 200)
        self.assertGreaterEqual(response.json["completion"]["percent"], 80)

        analysis = self.client.post("/api/analyze", json={"title": "QA Automation Engineer", "description": JOB_DESCRIPTION})
        self.assertEqual(analysis.status_code, 200)
        self.assertIn("Selenium", analysis.json["analysis"]["matched_skills"])
        self.assertIn("Playwright", analysis.json["analysis"]["missing_skills"])
        self.assertGreater(analysis.json["analysis"]["job_match_score"], 50)

        export = self.client.post("/api/resumes/generate", json={"title": "QA Automation Engineer", "description": JOB_DESCRIPTION})
        self.assertEqual(export.status_code, 200)
        self.assertEqual(export.mimetype, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        document = Document(BytesIO(export.data))
        document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Ava Tester", document_text)
        self.assertIn("CORE SKILLS", document_text)
        self.assertIn("Selenium", document_text)
        self.assertEqual(len(document.tables), 0)
        bullet = next(paragraph for paragraph in document.paragraphs if "Created Selenium" in paragraph.text)
        self.assertIn("<w:numPr>", bullet._p.xml)

    def test_job_decision_creates_application(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        created = self.client.post(
            "/api/jobs",
            json={"title": "QA Automation Engineer", "company": "Example Co", "description": JOB_DESCRIPTION},
        )
        self.assertEqual(created.status_code, 201)
        job_id = created.json["job"]["id"]
        decision = self.client.post(f"/api/jobs/{job_id}/decision", json={"status": "approved"})
        self.assertEqual(decision.status_code, 200)
        applications = self.client.get("/api/applications")
        self.assertEqual(len(applications.json["applications"]), 1)
        self.assertEqual(applications.json["applications"][0]["status"], "approved")

        rejected = self.client.post(f"/api/jobs/{job_id}/decision", json={"status": "rejected"})
        self.assertEqual(rejected.status_code, 200)
        self.assertEqual(self.client.get("/api/applications").json["applications"], [])

    def test_malformed_profile_does_not_erase_saved_data(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        response = self.client.put("/api/profile", data="{not valid json", content_type="application/json")
        self.assertEqual(response.status_code, 400)
        retained = self.client.get("/api/profile").json["profile"]
        self.assertEqual(retained["full_name"], "Ava Tester")

    def test_multiline_priority_and_additional_technology_gap(self) -> None:
        requirements = extract_requirements(
            """Required:\n- Selenium\n- C#\n- BrowserStack\nPreferred:\n- Playwright\n- Allure"""
        )
        priorities = {item["skill"]: item["priority"] for item in requirements}
        self.assertEqual(priorities["Selenium"], "required")
        self.assertEqual(priorities["C#"], "required")
        self.assertEqual(priorities["BrowserStack"], "required")
        self.assertEqual(priorities["Playwright"], "preferred")
        self.assertEqual(priorities["Allure"], "preferred")

    def test_current_form_export_overrides_saved_job_draft(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        created = self.client.post(
            "/api/jobs",
            json={"title": "Older QA role", "company": "Old Co", "description": JOB_DESCRIPTION},
        )
        job_id = created.json["job"]["id"]
        export = self.client.post(
            "/api/resumes/generate",
            json={
                "job_id": job_id,
                "title": "Current QA Automation Engineer",
                "company": "Current Co",
                "description": "Required: Selenium, API Testing, SQL. Preferred: Playwright.",
            },
        )
        self.assertEqual(export.status_code, 200)
        document = Document(BytesIO(export.data))
        self.assertIn("Current QA Automation Engineer", document.core_properties.title)

    def test_default_starter_is_editable_but_blocks_personalised_export(self) -> None:
        profile = self.client.get("/api/profile")
        self.assertEqual(profile.status_code, 200)
        self.assertTrue(profile.json["profile"]["is_starter_template"])
        self.assertEqual(profile.json["profile"]["full_name"], "Your Name")
        self.assertEqual(profile.json["completion"]["percent"], 0)

        starter_docx = self.client.get("/api/resumes/starter")
        self.assertEqual(starter_docx.status_code, 200)
        self.assertEqual(starter_docx.mimetype, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        blocked = self.client.post(
            "/api/resumes/generate",
            json={"title": "QA Engineer", "description": JOB_DESCRIPTION},
        )
        self.assertEqual(blocked.status_code, 422)

    def test_profile_json_round_trip_and_local_review(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        exported = self.client.get("/api/profile/export")
        self.assertEqual(exported.status_code, 200)
        saved = json.loads(exported.data.decode("utf-8"))
        self.assertEqual(saved["profile"]["full_name"], "Ava Tester")

        imported = self.client.post(
            "/api/profile/import",
            data={"file": (BytesIO(exported.data), "profile.json")},
            content_type="multipart/form-data",
        )
        self.assertEqual(imported.status_code, 200)
        self.assertEqual(imported.json["profile"]["email"], "ava@example.com")

        # Keep the suite deterministic even when a real local Ollama model is
        # installed: this path also verifies the transparent built-in fallback.
        with patch("local_ai.ollama_status", return_value={"available": False, "selected_installed": False}):
            review = self.client.post("/api/ai/review", json={"task": "proofread", "text": "I was responsible for teh QA work."})
        self.assertEqual(review.status_code, 200)
        self.assertIn("provider", review.json)
        self.assertIn("suggestions", review.json)

    def test_close_reopen_and_csv_job_import(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        created = self.client.post(
            "/api/jobs",
            json={"title": "QA Engineer", "company": "Example Co", "description": JOB_DESCRIPTION},
        )
        job_id = created.json["job"]["id"]
        self.client.post(f"/api/jobs/{job_id}/decision", json={"status": "approved"})
        closed = self.client.post(f"/api/jobs/{job_id}/close", json={})
        self.assertEqual(closed.status_code, 200)
        self.assertEqual(closed.json["job"]["status"], "closed")
        self.assertEqual(self.client.get("/api/applications").json["applications"], [])
        self.assertEqual(self.client.post(f"/api/jobs/{job_id}/reopen", json={}).json["job"]["status"], "new")

        csv_data = (
            "title,company,url,description,role_track\n"
            "API Test Engineer,Product Labs,https://example.com/jobs/api,"
            "Required API testing Postman SQL and regression testing for a SaaS product.,API Testing\n"
        ).encode("utf-8")
        imported = self.client.post(
            "/api/jobs/import",
            data={"file": (BytesIO(csv_data), "jobs.csv"), "source": "LinkedIn / user import"},
            content_type="multipart/form-data",
        )
        self.assertEqual(imported.status_code, 200)
        self.assertEqual(imported.json["added"], 1)
        sources = self.client.get("/api/job-sources")
        self.assertIn("Test Engineer", sources.json["role_tracks"])

    def test_qa_role_tracks_cover_manual_automation_sdet_and_api(self) -> None:
        cases = {
            "Manual QA Tester": "Manual QA",
            "QA Automation Engineer": "QA Automation",
            "Software Development Engineer in Test": "SDET",
            "API Test Engineer": "API Testing",
        }
        for title, expected_track in cases.items():
            classification = classify_qa_role(title, "Software product testing with API, regression, Agile and automation coverage.")
            self.assertIsNotNone(classification)
            self.assertEqual(classification[0], expected_track)
        self.assertIsNone(classify_qa_role("Software Engineer, Internal Systems", "Works with the QA team on a software platform."))
        self.assertIsNone(classify_qa_role("Staff Software Engineer", "Builds systems with the QA team."))
        self.assertIsNone(classify_qa_role("Fraud Operations Manager", "Partners with quality assurance and engineering teams."))

    def test_multi_source_discovery_merges_and_reports_public_sources(self) -> None:
        first = {
            "external_id": "remotive:1", "source": "Remotive", "source_url": "https://example.com/one",
            "title": "QA Automation Engineer", "company": "Example", "location": "Remote", "job_type": "Full-time",
            "description": "QA automation role with Selenium, API testing, SQL, and product delivery.", "role_track": "QA Automation",
            "quality_score": 70, "is_product_company": 0, "salary": "",
        }
        second = {
            **first, "external_id": "remote-ok:2", "source": "Remote OK", "source_url": "https://example.com/two",
            "title": "SDET", "role_track": "SDET", "quality_score": 85, "is_product_company": 1, "salary": "$100k",
        }
        with (
            patch("job_discovery.PRODUCT_COMPANY_BOARDS", []),
            patch("job_discovery.fetch_remotive", return_value=[first]),
            patch("job_discovery.fetch_remote_ok", return_value=[second]),
            patch("job_discovery.fetch_jobicy", return_value=[]),
            patch("job_discovery.fetch_himalayas", return_value=[]),
            patch("job_discovery.fetch_arbeitnow", return_value=[]),
            patch("job_discovery.fetch_the_muse", return_value=[]),
        ):
            jobs, report = discover_qa_jobs("qa test engineer")
        self.assertEqual([job["source"] for job in jobs], ["Remote OK", "Remotive"])
        self.assertEqual(
            {item["source"] for item in report},
            {"Remotive", "Remote OK", "Jobicy", "Himalayas", "Arbeitnow", "The Muse", "Google web job search"},
        )
        self.assertTrue(all(item["status"] == "ok" for item in report if item["source"] != "Google web job search"))
        self.assertEqual(next(item for item in report if item["source"] == "Google web job search")["status"], "not_configured")

    def test_discovery_cache_retains_india_results_without_daily_lock(self) -> None:
        self.client.put("/api/profile", json=SAMPLE_PROFILE)
        candidate = {
            "external_id": "mock:india-qa", "source": "Mock public feed", "source_url": "https://example.com/jobs/qa",
            "title": "QA Automation Engineer", "company": "Product Labs", "location": "Bengaluru, India", "job_type": "Full-time",
            "description": "QA automation role with Selenium, API testing, SQL, regression testing, and product delivery.",
            "role_track": "QA Automation", "quality_score": 88, "is_product_company": 1, "salary": "INR package disclosed",
        }
        report = [{"source": "Mock public feed", "status": "ok", "count": 1}]
        payload = {"query": "qa automation", "market": "India", "role_track": "All QA tracks", "include_product_boards": True}
        with patch("app.discover_qa_jobs", return_value=([candidate], report)) as discover:
            first = self.client.post("/api/jobs/discover", json=payload)
            second = self.client.post("/api/jobs/discover", json=payload)
        self.assertEqual(first.status_code, 200)
        self.assertFalse(first.json["cached"])
        self.assertEqual(first.json["reviewed"], 1)
        self.assertEqual(first.json["jobs"][0]["title"], "QA Automation Engineer")
        self.assertEqual(second.status_code, 200)
        self.assertTrue(second.json["cached"])
        self.assertEqual(discover.call_count, 1)
        self.assertEqual(self.client.get("/api/jobs?status=new").json["jobs"][0]["location"], "Bengaluru, India")

    def test_manual_application_keeps_hr_and_referral_details(self) -> None:
        created = self.client.post(
            "/api/applications",
            json={
                "title": "SDET", "company": "Product Labs", "application_kind": "Referral", "status": "approved",
                "location": "Pune, India", "role_track": "SDET", "description": "Referral opportunity for API and browser automation testing with Python and CI/CD.",
                "contact_name": "Asha Recruiter", "contact_email": "asha@example.com", "referral_name": "Ravi", "next_step": "Follow up Friday",
            },
        )
        self.assertEqual(created.status_code, 201)
        application = self.client.get("/api/applications").json["applications"][0]
        self.assertEqual(application["application_kind"], "Referral")
        self.assertEqual(application["contact_name"], "Asha Recruiter")
        self.assertEqual(application["referral_name"], "Ravi")

    def test_india_filter_does_not_accept_a_foreign_role_with_india_in_the_description(self) -> None:
        foreign = {"location": "USA", "description": "Our global teams include colleagues in India.", "title": "QA Engineer"}
        india = {"location": "Bengaluru, India", "description": "QA automation role.", "title": "QA Engineer"}
        remote_india = {"location": "Remote", "description": "Remote work is available for candidates based in India.", "title": "QA Engineer"}
        self.assertFalse(matches_market(foreign, "India"))
        self.assertTrue(matches_market(india, "India"))
        self.assertTrue(matches_market(remote_india, "India"))

    def test_chat_resume_request_returns_real_templates_and_a_truth_first_draft(self) -> None:
        reply = _resume_chat_reply("Give me templates and draft a resume", SAMPLE_PROFILE)
        self.assertIsNotNone(reply)
        self.assertIn("Classic ATS", reply)
        self.assertIn("Compact QA", reply)
        self.assertIn("Ava Tester", reply)
        self.assertIn("Created Selenium regression tests", reply)


if __name__ == "__main__":
    unittest.main(verbosity=2)
