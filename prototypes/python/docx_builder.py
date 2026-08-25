"""ATS-safe DOCX creation using real Word paragraphs and list bullets."""

from __future__ import annotations

from io import BytesIO
import re
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(21, 32, 51)
MUTED = RGBColor(77, 91, 112)
BLUE = RGBColor(29, 78, 216)
TEAL = RGBColor(15, 126, 113)

RESUME_LAYOUTS = {
    "classic": {"label": "Classic ATS", "font": "Aptos", "body": 10.5, "name": 20, "margin": 0.70, "accent": INK, "centered": True},
    "compact": {"label": "Compact QA", "font": "Aptos", "body": 9.6, "name": 18, "margin": 0.55, "accent": TEAL, "centered": False},
    "modern": {"label": "Modern single-column", "font": "Aptos", "body": 10.2, "name": 21, "margin": 0.68, "accent": BLUE, "centered": False},
}


def available_resume_layouts() -> list[dict[str, str]]:
    return [{"id": key, "label": value["label"]} for key, value in RESUME_LAYOUTS.items()]


def safe_text(value: Any) -> str:
    return str(value or "").replace("\x00", "").strip()


def add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CareerCraft Section")
    paragraph.add_run(text.upper())


def add_bullet(document: Document, text: str) -> None:
    if not safe_text(text):
        return
    paragraph = document.add_paragraph(style="CareerCraft Bullet")
    # A visible indentation is not enough: emit actual Word list numbering so
    # both Word and ATS text extractors identify these as real bullets.
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:numPr"):
            p_pr.remove(child)
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    # python-docx's default template maps numId=1 to a bullet definition.
    num_id.set(qn("w:val"), "1")
    num_pr.extend((level, num_id))
    p_pr.append(num_pr)
    paragraph.add_run(safe_text(text))


def build_resume_document(
    resume: dict[str, Any], target_title: str = "", company: str = "", layout: str = "classic"
) -> BytesIO:
    """Return a one-column Word document designed for reliable text extraction."""
    layout_id = layout if layout in RESUME_LAYOUTS else "classic"
    config = RESUME_LAYOUTS[layout_id]
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(config["margin"])
    section.right_margin = Inches(config["margin"])

    normal = document.styles["Normal"]
    normal.font.name = config["font"]
    normal.font.size = Pt(config["body"])
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.06

    for style_name, font_size in (("CareerCraft Section", config["body"] + 0.6), ("CareerCraft Bullet", config["body"])):
        if style_name not in document.styles:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = document.styles[style_name]
        style.font.name = config["font"]
        style.font.size = Pt(font_size)
        style.font.color.rgb = config["accent"]

    heading_style = document.styles["CareerCraft Section"]
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(9)
    heading_style.paragraph_format.space_after = Pt(3)

    bullet_style = document.styles["CareerCraft Bullet"]
    bullet_style.paragraph_format.left_indent = Inches(0.19)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.15)
    bullet_style.paragraph_format.space_after = Pt(1.5)

    name = safe_text(resume.get("full_name")) or "Your Name"
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if config["centered"] else WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_paragraph.add_run(name)
    name_run.bold = True
    name_run.font.name = "Aptos Display"
    name_run.font.size = Pt(config["name"])
    name_run.font.color.rgb = config["accent"]
    name_paragraph.paragraph_format.space_after = Pt(2)

    headline = safe_text(resume.get("headline")) or safe_text(target_title)
    if headline:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if config["centered"] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(headline)
        run.bold = True
        run.font.size = Pt(config["body"])
        run.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(3)

    contact = [
        safe_text(resume.get("location")),
        safe_text(resume.get("phone")),
        safe_text(resume.get("email")),
        safe_text(resume.get("linkedin_url")),
        safe_text(resume.get("portfolio_url")),
    ]
    contact = [part for part in contact if part]
    if contact:
        p = document.add_paragraph(" | ".join(contact))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if config["centered"] else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(7)
        for run in p.runs:
            run.font.size = Pt(max(8.8, config["body"] - 1))
            run.font.color.rgb = MUTED

    summary = safe_text(resume.get("summary"))
    if summary:
        add_section_heading(document, "Professional Summary")
        document.add_paragraph(summary)

    skills = [safe_text(skill) for skill in resume.get("skills") or [] if safe_text(skill)]
    if skills:
        add_section_heading(document, "Core Skills")
        paragraph = document.add_paragraph(", ".join(skills))
        paragraph.paragraph_format.space_after = Pt(4)

    experience = resume.get("experience") or []
    if experience:
        add_section_heading(document, "Professional Experience")
        for entry in experience:
            position = safe_text(entry.get("title"))
            company_name = safe_text(entry.get("company"))
            location = safe_text(entry.get("location"))
            dates = " – ".join(
                part for part in (safe_text(entry.get("start_date")), "Present" if entry.get("current") else safe_text(entry.get("end_date"))) if part
            )
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            main = " | ".join(part for part in (position, company_name) if part)
            if main:
                run = p.add_run(main)
                run.bold = True
                run.font.size = Pt(10.7)
            details = " | ".join(part for part in (location, dates) if part)
            if details:
                details_run = p.add_run((" — " if main else "") + details)
                details_run.font.color.rgb = MUTED
                details_run.font.size = Pt(9.8)
            for bullet in entry.get("bullets") or []:
                add_bullet(document, bullet)

    projects = resume.get("projects") or []
    if projects:
        add_section_heading(document, "Selected Projects")
        for project in projects:
            name_value = safe_text(project.get("name"))
            description = safe_text(project.get("description"))
            p = document.add_paragraph()
            if name_value:
                run = p.add_run(name_value)
                run.bold = True
            if description:
                p.add_run((" — " if name_value else "") + description)
            for bullet in project.get("bullets") or []:
                add_bullet(document, bullet)

    education = resume.get("education") or []
    if education:
        add_section_heading(document, "Education")
        for entry in education:
            degree = safe_text(entry.get("degree"))
            school = safe_text(entry.get("school"))
            location = safe_text(entry.get("location"))
            graduation = safe_text(entry.get("graduation"))
            p = document.add_paragraph()
            text = " | ".join(part for part in (degree, school) if part)
            if text:
                p.add_run(text).bold = True
            details = " | ".join(part for part in (location, graduation) if part)
            if details:
                details_run = p.add_run((" — " if text else "") + details)
                details_run.font.color.rgb = MUTED

    certifications = [safe_text(item) for item in resume.get("certifications") or [] if safe_text(item)]
    if certifications:
        add_section_heading(document, "Certifications")
        document.add_paragraph(" | ".join(certifications))

    document.core_properties.title = f"{name} — {target_title or 'Resume'}"
    document.core_properties.subject = "ATS-friendly resume"
    document.core_properties.author = name
    document.core_properties.comments = f"CareerCraft {config['label']} layout"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def resume_filename(profile: dict[str, Any], target_title: str = "") -> str:
    name = safe_text(profile.get("full_name")) or "Resume"
    title = safe_text(target_title) or "QA Resume"
    value = re.sub(r"[^A-Za-z0-9]+", "_", f"{name}_{title}").strip("_")
    return f"{value[:100] or 'Resume'}.docx"
